import io
import os
import sys
import tempfile
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import parse_qs, urlparse
from pypdf import PdfReader
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def pdf_to_docx_bytes(pdf_bytes):
    """Converts raw PDF bytes into a clean, editable Word (.docx) document byte stream."""
    pdf_file = io.BytesIO(pdf_bytes)
    reader = PdfReader(pdf_file)
    doc = docx.Document()

    # Set standard 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if not text or not text.strip():
            continue

        lines = text.splitlines()
        for line in lines:
            trimmed = line.strip()
            if not trimmed:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)

            # Heading vs Body Paragraph heuristic
            if len(trimmed) < 60 and not trimmed.endswith('.'):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(trimmed)
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(30, 41, 59)
            elif trimmed.startswith(('•', '-', '*', '1.', '2.', '3.')):
                run = p.add_run(trimmed)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(51, 65, 85)
            else:
                run = p.add_run(trimmed)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(15, 23, 42)

    output_stream = io.BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()

class PDFConverterRequestHandler(BaseHTTPRequestHandler):
    def _send_cors_headers(self):
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        self.send_header('Cache-Control', 'no-cache, no-store, must-revalidate')

    def do_OPTIONS(self):
        self.send_response(200)
        self._send_cors_headers()
        self.end_headers()

    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path in ('/health', '/v1/health'):
            self.send_response(200)
            self._send_cors_headers()
            self.send_header('Content-Type', 'text/plain')
            self.end_headers()
            self.wfile.write(b"OK")
        else:
            self.send_response(404)
            self._send_cors_headers()
            self.end_headers()

    def do_POST(self):
        parsed = urlparse(self.path)
        if parsed.path in ('/convert/pdf-to-docx', '/v1/convert'):
            try:
                content_type = self.headers.get('Content-Type', '')
                content_length = int(self.headers.get('Content-Length', 0))

                body = self.rfile.read(content_length)

                # Extract PDF file bytes from multipart form-data or raw body
                pdf_bytes = None
                if 'boundary=' in content_type:
                    boundary = content_type.split('boundary=')[1].encode('ascii')
                    parts = body.split(b'--' + boundary)
                    for part in parts:
                        if b'filename=' in part or b'name="file"' in part:
                            header_end = part.find(b'\r\n\r\n')
                            if header_end != -1:
                                pdf_bytes = part[header_end + 4:].rstrip(b'\r\n--')
                                break
                else:
                    pdf_bytes = body

                if not pdf_bytes or len(pdf_bytes) < 10:
                    self.send_response(400)
                    self._send_cors_headers()
                    self.send_header('Content-Type', 'application/json')
                    self.end_headers()
                    self.wfile.write(b'{"error": "No valid PDF file received."}')
                    return

                docx_bytes = pdf_to_docx_bytes(pdf_bytes)

                self.send_response(200)
                self._send_cors_headers()
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', 'attachment; filename="converted.docx"')
                self.send_header('Content-Length', str(len(docx_bytes)))
                self.end_headers()
                self.wfile.write(docx_bytes)

            except Exception as e:
                self.send_response(500)
                self._send_cors_headers()
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                error_msg = f'{{"error": "Conversion Exception: {str(e)}"}}'
                self.wfile.write(error_msg.encode('utf-8'))
        else:
            self.send_response(404)
            self._send_cors_headers()
            self.end_headers()

def run_server(port=8080):
    server_address = ('0.0.0.0', port)
    httpd = HTTPServer(server_address, PDFConverterRequestHandler)
    print(f"🚀 PDF to Word Converter Service running on http://localhost:{port}")
    httpd.serve_forever()

if __name__ == '__main__':
    run_server(8080)
