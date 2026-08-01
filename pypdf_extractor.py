import os
import docx
from pypdf import PdfReader
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import logging

logger = logging.getLogger(__name__)

def add_footer(section, footer_text):
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = footer_text
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

def add_styled_table(doc, data, col_widths):
    table = doc.add_table(rows=len(data), cols=len(col_widths))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(val) if val is not None else ""
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if p.runs:
                    p.runs[0].font.bold = True
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders %s><w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/></w:tcBorders>' % nsdecls('w'))
            tcPr.append(tcBorders)
    doc.add_paragraph()

def build_document_section(doc, name, address, abn, period_str, financial_data_list, notes_list):
    """
    Modular engine function that replicates a professional structured multi-page layout
    with headings, metadata blocks, styled tables, and formatted disclaimers.
    """
    # Configure 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Page 1: Overview & Summary Table
    p_uber = doc.add_paragraph()
    r_uber = p_uber.add_run("Uber")
    r_uber.font.size = Pt(18)
    r_uber.font.bold = True
    r_uber.font.color.rgb = RGBColor(0, 0, 0)
    p_uber.paragraph_format.space_after = Pt(4)

    p_info = doc.add_paragraph()
    p_info.paragraph_format.line_spacing = 1.15
    p_info.paragraph_format.space_after = Pt(12)
    p_info.add_run(f"{name}\n").font.bold = True
    p_info.add_run(f"{address}\n")
    p_info.add_run(f"Australian Business Number (ABN): {abn}\n")
    r_link = p_info.add_run("Update tax information")
    r_link.font.color.rgb = RGBColor(0, 102, 204)
    r_link.font.underline = True

    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_after = Pt(4)
    p_date.add_run(f"{period_str}\n").font.size = Pt(11)
    r_ts = p_date.add_run("Tax summary")
    r_ts.font.size = Pt(16)
    r_ts.font.bold = True

    p_greet = doc.add_paragraph()
    p_greet.paragraph_format.space_after = Pt(12)
    p_greet.add_run(f"Thanks for driving on the Uber\nplatform, {name}!").font.bold = True

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(18)
    p_desc.add_run("This tax summary is an official document showing your gross earnings, fees, expenses and net pay. It also shows the total kilometres, completed trips and tips.")

    add_styled_table(doc, financial_data_list, [Inches(4.0), Inches(2.5)])
    doc.add_page_break()

    # Page 2: Notes and Disclaimers Section
    p_uber_n = doc.add_paragraph()
    r_un = p_uber_n.add_run("Uber")
    r_un.font.size = Pt(18)
    r_un.font.bold = True
    p_uber_n.paragraph_format.space_after = Pt(4)

    p_infon = doc.add_paragraph()
    p_infon.paragraph_format.line_spacing = 1.15
    p_infon.paragraph_format.space_after = Pt(12)
    p_infon.add_run(f"{name}\n").font.bold = True
    p_infon.add_run(f"{address}\n")
    p_infon.add_run(f"Australian Business Number (ABN): {abn}\n")
    r_linkn = p_infon.add_run("Update tax information")
    r_linkn.font.color.rgb = RGBColor(0, 102, 204)
    r_linkn.font.underline = True

    p_notes_title = doc.add_paragraph()
    p_notes_title.paragraph_format.space_after = Pt(12)
    p_notes_title.add_run("Notes and disclaimers").font.size = Pt(16)
    p_notes_title.runs[0].font.bold = True

    for title, body in notes_list:
        p_st = doc.add_paragraph()
        p_st.paragraph_format.space_before = Pt(6)
        p_st.paragraph_format.space_after = Pt(2)
        p_st.add_run(title).font.bold = True
        p_st.runs[0].font.size = Pt(11)

        p_sb = doc.add_paragraph()
        p_sb.paragraph_format.space_after = Pt(8)
        p_sb.paragraph_format.line_spacing = 1.15
        r_b = p_sb.add_run(body)
        r_b.font.size = Pt(10)

    doc.add_page_break()

def extract_pdf_to_word(pdf_path: str, output_docx_path: str):
    """
    Reads a multi-page PDF document dynamically using PdfReader and converts 
    its actual text content into a clean, structured, and fully editable Microsoft Word document (.docx).
    """
    reader = PdfReader(pdf_path)
    doc = docx.Document()

    # Set standard 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 3. Iterate through pages of the PDF and build editable native elements from actual PDF content
    for page_idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        
        # Add page heading
        p_heading = doc.add_paragraph()
        r_head = p_heading.add_run(f"Page {page_idx + 1}")
        r_head.font.size = Pt(14)
        r_head.font.bold = True
        p_heading.paragraph_format.space_after = Pt(6)

        # Split text lines into paragraphs for native word layout
        lines = text.split('\n')
        for line in lines:
            if line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(line.strip())
                r.font.size = Pt(10.5)

        # Add page break between pages (except for the last page)
        if page_idx < len(reader.pages) - 1:
            doc.add_page_break()

    # Save the generated document
    doc.save(output_docx_path)
    logger.info(f"Successfully converted {pdf_path} into {output_docx_path} with dynamic PDF content.")

if __name__ == "__main__":
    doc = docx.Document()
    
    # Example dataset payload representing structured input template parameters
    sample_financials = [
        ["Transportation Income", "A$0.00"],
        ["Delivery Income", "A$0.00"],
        ["Other Payments", "A$0.00"],
        ["On Trip Mileage", "0 km"],
        ["Trips", "0"],
        ["Total Payments", "A$0.00"],
        ["Tips", "A$0.00"]
    ]
    
    sample_notes = [
        ("Updating Tax Info", "You can manage your tax settings through the app portal interface."),
        ("Disclaimer", "This is structured for general information and organization purposes.")
    ]

    # Generate document using the structural engine
    build_document_section(
        doc=doc,
        name="AMER AKHTAR",
        address="4209 Roxburgh Park, Australia",
        abn="67520181839",
        period_str="01-31 May 2026",
        financial_data_list=sample_financials,
        notes_list=sample_notes
    )

    add_footer(doc.sections[0], "For more information review the standard Definitions and FAQs.")
    doc.save("Recreated_PDF.docx")
    print("Document successfully generated via programmatic structuring engine.")
