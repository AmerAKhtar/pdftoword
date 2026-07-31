"""
Alignment Analyzer & Layout Alignment Fixer Module
Compares PDF layout geometry (PyMuPDF fitz) with generated Word (.docx) documents (python-docx)
and automatically corrects margins, paragraph alignments, indentations, and table positions.
"""

import logging
import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, Length

logger = logging.getLogger(__name__)


def analyze_pdf_layout(pdf_path: str) -> dict:
    """
    Analyzes the source PDF to extract page dimensions, margins, and text block alignments.
    """
    doc = fitz.open(pdf_path)
    page_data = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        rect = page.rect  # (0, 0, width, height)
        width, height = rect.width, rect.height

        blocks = []
        d = page.get_text("dict")
        
        min_x = width
        max_x = 0
        min_y = height
        max_y = 0

        for b in d.get("blocks", []):
            if b.get("type") == 0:  # Text block
                bbox = b["bbox"]  # (x0, y0, x1, y1)
                x0, y0, x1, y1 = bbox
                
                min_x = min(min_x, x0)
                max_x = max(max_x, x1)
                min_y = min(min_y, y0)
                max_y = max(max_y, y1)

                lines_text = []
                for ln in b.get("lines", []):
                    line_text = "".join(sp.get("text", "") for sp in ln.get("spans", [])).strip()
                    if line_text:
                        lines_text.append(line_text)

                block_text = " ".join(lines_text).strip()
                if not block_text:
                    continue

                # Calculate block alignment characteristics
                block_width = x1 - x0
                center_x = (x0 + x1) / 2.0
                dist_from_center = abs(center_x - (width / 2.0))

                alignment = "LEFT"
                if dist_from_center < (width * 0.06) and block_width < (width * 0.85):
                    alignment = "CENTER"
                elif (width - x1) < (x0 * 0.5) and (x0 > width * 0.5):
                    alignment = "RIGHT"
                elif block_width > (width * 0.75):
                    alignment = "JUSTIFY"

                blocks.append({
                    "bbox": bbox,
                    "width": block_width,
                    "center_x": center_x,
                    "alignment": alignment,
                    "text": block_text
                })

        # Calculate estimated margins
        margin_left = min_x if min_x < width else 72.0
        margin_right = (width - max_x) if max_x > 0 else 72.0
        margin_top = min_y if min_y < height else 72.0
        margin_bottom = (height - max_y) if max_y > 0 else 72.0

        page_data.append({
            "page_num": page_num,
            "width": width,
            "height": height,
            "margins": {
                "left": margin_left,
                "right": margin_right,
                "top": margin_top,
                "bottom": margin_bottom
            },
            "blocks": blocks
        })

    doc.close()
    return {"pages": page_data}


def align_docx_with_pdf(pdf_path: str, docx_path: str, output_path: str) -> str:
    """
    Compares PDF block positioning with DOCX paragraph/table structures and fixes alignments.
    """
    logger.info(f"Analyzing PDF layout alignment from {pdf_path}...")
    pdf_info = analyze_pdf_layout(pdf_path)

    doc = Document(docx_path)
    pages_info = pdf_info.get("pages", [])

    if not pages_info:
        doc.save(output_path)
        return output_path

    # Page 1 geometry reference
    ref_page = pages_info[0]
    ref_margins = ref_page["margins"]
    pdf_width = ref_page["width"]
    pdf_height = ref_page["height"]

    # 1. Update DOCX section margins to match PDF geometry
    for section in doc.sections:
        try:
            section.page_width = Pt(pdf_width)
            section.page_height = Pt(pdf_height)
            
            # Apply bounds checking for margins (0.5 in to 1.5 in)
            left_pt = max(36.0, min(ref_margins["left"], 108.0))
            right_pt = max(36.0, min(ref_margins["right"], 108.0))
            top_pt = max(36.0, min(ref_margins["top"], 108.0))
            bottom_pt = max(36.0, min(ref_margins["bottom"], 108.0))

            section.left_margin = Pt(left_pt)
            section.right_margin = Pt(right_pt)
            section.top_margin = Pt(top_pt)
            section.bottom_margin = Pt(bottom_pt)
        except Exception as e:
            logger.warning(f"Could not adjust section margins: {e}")

    # Index PDF blocks for alignment comparison
    pdf_blocks = []
    for page in pages_info:
        pdf_blocks.extend(page.get("blocks", []))

    # 2. Fix paragraph alignments based on PDF block matching
    pdf_idx = 0
    num_pdf_blocks = len(pdf_blocks)

    for paragraph in doc.paragraphs:
        p_text = paragraph.text.strip()
        if not p_text:
            continue

        # Find matching block in PDF
        matched_block = None
        if pdf_idx < num_pdf_blocks:
            pdf_block = pdf_blocks[pdf_idx]
            # Simple text match heuristic
            if p_text[:15] in pdf_block["text"] or pdf_block["text"][:15] in p_text:
                matched_block = pdf_block
                pdf_idx += 1
            else:
                # Look ahead 3 blocks
                for ahead in range(1, 4):
                    if pdf_idx + ahead < num_pdf_blocks:
                        cand = pdf_blocks[pdf_idx + ahead]
                        if p_text[:15] in cand["text"] or cand["text"][:15] in p_text:
                            matched_block = cand
                            pdf_idx += ahead + 1
                            break

        if matched_block:
            target_align = matched_block["alignment"]
            if target_align == "CENTER":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif target_align == "RIGHT":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif target_align == "JUSTIFY":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif target_align == "LEFT":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Fallback heuristic: check paragraph length and text
            if len(p_text) < 40 and paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                # Short standalone titles/headers that look centered
                pass

    # 3. Fix tables alignment (Center all tables on page)
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for cell_p in cell.paragraphs:
                    # Keep cell text left-aligned or centered cleanly
                    if not cell_p.alignment:
                        cell_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(output_path)
    logger.info(f"Alignment optimization completed -> {output_path}")
    return output_path


if __name__ == "__main__":
    import sys
    if len(sys.argv) >= 3:
        src_pdf = sys.argv[1]
        src_docx = sys.argv[2]
        out_docx = sys.argv[3] if len(sys.argv) > 3 else "aligned_output.docx"
        align_docx_with_pdf(src_pdf, src_docx, out_docx)
        print(f"Alignment fixed: {out_docx}")
    else:
        print("Usage: python fix_alignment.py <input.pdf> <input.docx> [output.docx]")
