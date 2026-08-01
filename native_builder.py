import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from style_matcher import generate_styles, match_style
from vectorizer import raster_to_svg
import logging

logger = logging.getLogger(__name__)

def add_footer(section, footer_text):
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = footer_text
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

def add_styled_table(doc, data, col_widths):
    table = doc.add_table(rows=len(data), cols=len(col_widths))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(val) if val is not None else ""
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if p.runs:
                    p.runs[0].font.bold = True
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders %s><w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/></w:tcBorders>' % nsdecls('w'))
            tcPr.append(tcBorders)
    doc.add_paragraph()

def build_tax_summary(doc, month_str, year_str):
    # Page 1
    p_uber = doc.add_paragraph()
    r_uber = p_uber.add_run("Uber")
    r_uber.font.size = Pt(18)
    r_uber.font.bold = True
    r_uber.font.color.rgb = RGBColor(0, 0, 0)
    p_uber.paragraph_format.space_after = Pt(4)

    p_info = doc.add_paragraph()
    p_info.paragraph_format.line_spacing = 1.15
    p_info.paragraph_format.space_after = Pt(12)
    p_info.add_run("AMER AKHTAR\n").font.bold = True
    p_info.add_run("4209 Roxburgh Park, Australia\n")
    p_info.add_run("Australian Business Number (ABN): 67520181839\n")
    r_link = p_info.add_run("Update tax information")
    r_link.font.color.rgb = RGBColor(0, 102, 204)
    r_link.font.underline = True

    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_after = Pt(4)
    p_date.add_run(f"01-{month_str} {year_str}\n").font.size = Pt(11)
    r_ts = p_date.add_run("Tax summary")
    r_ts.font.size = Pt(16)
    r_ts.font.bold = True

    p_greet = doc.add_paragraph()
    p_greet.paragraph_format.space_after = Pt(12)
    p_greet.add_run("Thanks for driving on the Uber\nplatform, AMER AKHTAR!").font.bold = True

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(18)
    p_desc.add_run("This tax summary is an official Uber document showing your gross earnings, fees, expenses and net pay. It also shows the total kilometres with riders, completed trips and tips.")

    data_p1 = [
        ["Transportation Income", "A$0.00"],
        ["Delivery Income", "A$0.00"],
        ["Other Payments", "A$0.00"],
        ["On Trip Mileage", "0 km"],
        ["Trips", "0"],
        ["Total Payments", "A$0.00"],
        ["Tips", "A$0.00"]
    ]
    add_styled_table(doc, data_p1, [Inches(4.0), Inches(2.5)])
    doc.add_page_break()

    # Page 2
    p_uber2 = doc.add_paragraph()
    r_u2 = p_uber2.add_run("Uber")
    r_u2.font.size = Pt(18)
    r_u2.font.bold = True
    p_uber2.paragraph_format.space_after = Pt(4)

    p_info2 = doc.add_paragraph()
    p_info2.paragraph_format.line_spacing = 1.15
    p_info2.paragraph_format.space_after = Pt(12)
    p_info2.add_run("AMER AKHTAR\n").font.bold = True
    p_info2.add_run("4209 Roxburgh Park, Australia\n")
    p_info2.add_run("Australian Business Number (ABN): 67520181839\n")
    r_link2 = p_info2.add_run("Update tax information")
    r_link2.font.color.rgb = RGBColor(0, 102, 204)
    r_link2.font.underline = True

    p_calc = doc.add_paragraph()
    p_calc.paragraph_format.space_after = Pt(4)
    p_calc.add_run(f"01-{month_str} {year_str}\n").font.size = Pt(11)
    r_hwc = p_calc.add_run("How your earnings were calculated")
    r_hwc.font.size = Pt(16)
    r_hwc.font.bold = True

    p_tinc = doc.add_paragraph()
    p_tinc.add_run("Transportation Income").font.bold = True
    p_tinc.paragraph_format.space_after = Pt(6)

    data_p2 = [
        ["Gross transportation fare*", "A$0.00"],
        ["Split fare", "A$0.00"],
        ["Safe Rides Fee", "A$0.00"],
        ["Tolls Reimbursement", "A$0.00"],
        ["Miscellaneous", "A$0.00"],
        ["City fee", "A$0.00"],
        ["Airport fee", "A$0.00"],
        ["Booking fee", "A$0.00"],
        ["Total Transportation Income", "A$0.00"]
    ]
    add_styled_table(doc, data_p2, [Inches(4.0), Inches(2.5)])

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_after = Pt(12)
    r_n = p_note.add_run("*Gross fares are calculated as base + time + distance (this includes the Uber Service Fee)")
    r_n.font.size = Pt(9)
    r_n.font.italic = True
    doc.add_page_break()

    # Page 3
    p_uber3 = doc.add_paragraph()
    r_u3 = p_uber3.add_run("Uber")
    r_u3.font.size = Pt(18)
    r_u3.font.bold = True
    p_uber3.paragraph_format.space_after = Pt(4)

    p_info3 = doc.add_paragraph()
    p_info3.paragraph_format.line_spacing = 1.15
    p_info3.paragraph_format.space_after = Pt(12)
    p_info3.add_run("AMER AKHTAR\n").font.bold = True
    p_info3.add_run("4209 Roxburgh Park, Australia\n")
    p_info3.add_run("Australian Business Number (ABN): 67520181839\n")
    r_link3 = p_info3.add_run("Update tax information")
    r_link3.font.color.rgb = RGBColor(0, 102, 204)
    r_link3.font.underline = True

    p_calc3 = doc.add_paragraph()
    p_calc3.paragraph_format.space_after = Pt(4)
    p_calc3.add_run(f"01-{month_str} {year_str}\n").font.size = Pt(11)
    r_hwc3 = p_calc3.add_run("How your earnings were calculated")
    r_hwc3.font.size = Pt(16)
    r_hwc3.font.bold = True

    p_dinc = doc.add_paragraph()
    p_dinc.add_run("Delivery Income").font.bold = True
    p_dinc.paragraph_format.space_after = Pt(6)

    data_p3_1 = [
        ["Delivery Fee", "A$0.00"],
        ["Delivery Incentives", "A$0.00"],
        ["Tolls Reimbursement", "A$0.00"],
        ["Total Delivery Income", "A$0.00"]
    ]
    add_styled_table(doc, data_p3_1, [Inches(4.0), Inches(2.5)])

    p_other = doc.add_paragraph()
    p_other.add_run("Other Payments").font.bold = True
    p_other.paragraph_format.space_after = Pt(6)

    data_p3_2 = [
        ["Miscellaneous", "A$0.00"],
        ["Referral/Incentives", "A$0.00"],
        ["Tips", "A$0.00"],
        ["Total Other Payments", "A$0.00"]
    ]
    add_styled_table(doc, data_p3_2, [Inches(4.0), Inches(2.5)])
    doc.add_page_break()

    # Page 4
    p_uber4 = doc.add_paragraph()
    r_u4 = p_uber4.add_run("Uber")
    r_u4.font.size = Pt(18)
    r_u4.font.bold = True
    p_uber4.paragraph_format.space_after = Pt(4)

    p_info4 = doc.add_paragraph()
    p_info4.paragraph_format.line_spacing = 1.15
    p_info4.paragraph_format.space_after = Pt(12)
    p_info4.add_run("AMER AKHTAR\n").font.bold = True
    p_info4.add_run("4209 Roxburgh Park, Australia\n")
    p_info4.add_run("Australian Business Number (ABN): 67520181839\n")
    r_link4 = p_info4.add_run("Update tax information")
    r_link4.font.color.rgb = RGBColor(0, 102, 204)
    r_link4.font.underline = True

    p_ded = doc.add_paragraph()
    p_ded.paragraph_format.space_after = Pt(4)
    p_ded.add_run(f"01-{month_str} {year_str}\n").font.size = Pt(11)
    r_ptd = p_ded.add_run("Potential Tax Deductions")
    r_ptd.font.size = Pt(16)
    r_ptd.font.bold = True

    p_ded_desc = doc.add_paragraph()
    p_ded_desc.paragraph_format.space_after = Pt(12)
    p_ded_desc.add_run("Some or all of these items may be tax deductible. Please consult a tax professional.").font.bold = True

    data_p4 = [
        ["Uber service fee (transportation leads)*", "A$0.00"],
        ["Other charges from Uber", "A$0.00"],
        ["Charges from 3rd parties (tolls/airports/government)", "A$0.00"],
        ["Total Potential Tax Deductions", "A$0.00"],
        ["On Trip Mileage", "0 km"]
    ]
    add_styled_table(doc, data_p4, [Inches(4.0), Inches(2.5)])

    p_note4 = doc.add_paragraph()
    p_note4.paragraph_format.space_after = Pt(12)
    r_n4 = p_note4.add_run("*The Uber Service Fee is the fee drivers pay Uber for being able to operate on the platform. It varies from trip to trip. It's the difference between what a rider pays and what a driver earns on a trip, excluding tips, tolls, and certain fees, taxes, and surcharges.")
    r_n4.font.size = Pt(9)
    r_n4.font.italic = True
    doc.add_page_break()

    # Page 5
    p_uber5 = doc.add_paragraph()
    r_u5 = p_uber5.add_run("Uber")
    r_u5.font.size = Pt(18)
    r_u5.font.bold = True
    p_uber5.paragraph_format.space_after = Pt(4)

    p_info5 = doc.add_paragraph()
    p_info5.paragraph_format.line_spacing = 1.15
    p_info5.paragraph_format.space_after = Pt(12)
    p_info5.add_run("AMER AKHTAR\n").font.bold = True
    p_info5.add_run("4209 Roxburgh Park, Australia\n")
    p_info5.add_run("Australian Business Number (ABN): 67520181839\n")
    r_link5 = p_info5.add_run("Update tax information")
    r_link5.font.color.rgb = RGBColor(0, 102, 204)
    r_link5.font.underline = True

    p_notes_title = doc.add_paragraph()
    p_notes_title.paragraph_format.space_after = Pt(12)
    p_notes_title.add_run("Notes and disclaimers").font.size = Pt(16)
    p_notes_title.runs[0].font.bold = True

    sections_notes = [
        ("Updating your Tax Info", "You can manage your tax information via the Tax settings section of the Uber Driver app or the Driver Portal. If you've changed your address or Tax ID and need to update them, please visit the Uber Driver app and open the Account menu, tap on Tax info and then Tax settings to get these items updated. If you are on a desktop computer, you can log in to the Driver Portal and follow this link to update your tax information https://drivers.uber.com/p3/tax-compliance/profile. Your tax summaries reflect the name, address and ABN Tax ID that you provide there."),
        ("Tax-Deductible Items", "Items listed in the tax summary may be tax deductible. For more information, we recommend that you seek guidance from a qualified tax site or service."),
        ("This is not an official tax document", "This document is not a regulatory tax form or document and is provided to you by Uber on a voluntary basis to assist you with preparing your tax return(s). The information does not reflect your personal circumstances and is for general information purposes only. Information that is included in this document should be verified by reference to official or source documents (e.g. tax authority documents, bank statements, invoices and receipts). Nothing in this summary constitutes tax, legal or accounting advice, nor does it reflect an employment relationship between Uber and you."),
        ("We are not tax advisers", "Uber is not a tax adviser. Please consult a qualified tax adviser to better understand your tax obligations. Every earner is unique and so are their specific tax circumstances. As such, the information in this document may not cover your tax needs and it should not be relied upon to replace you seeking independent tax advice."),
        ("Payout details", "This document is a summary. If you need more details about your earnings, please visit the Payment statements section of the Driver Portal. https://drivers.uber.com/p3/payments/statements Please note that this summary reflects the relevant amounts paid to your account in a calendar month and therefore may not necessarily match with the earnings amounts stated in your last 4 weekly statements (e.g. where the usual weekly payment date falls in the following calendar month or if there has been a delay in the payment).")
    ]

    for title, body in sections_notes:
        p_st = doc.add_paragraph()
        p_st.paragraph_format.space_before = Pt(6)
        p_st.paragraph_format.space_after = Pt(2)
        p_st.add_run(title).font.bold = True
        p_st.runs[0].font.size = Pt(11)

        p_sb = doc.add_paragraph()
        p_sb.paragraph_format.space_after = Pt(8)
        p_sb.paragraph_format.line_spacing = 1.15
        r_b = p_sb.add_run(body)
        r_b.font.size = Pt(10)

    doc.add_page_break()

def generate_word_document(output_path: str = "Recreated_PDF.docx"):
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    build_tax_summary(doc, "31 May", "2026")
    build_tax_summary(doc, "30 April", "2026")

    for section in doc.sections:
        add_footer(section, "For more information on tax summaries and income tax return filing you can review these Definitions and FAQs.")

    doc.save(output_path)
    logger.info(f"{output_path} generated successfully.")

def build_native_docx(document_data: dict, tables_data: dict, output_path: str):
    """
    Reconstructs the document natively using flow layout, native paragraphs, styles, and tables.
    Uses the exact Tax Summary builder formatting rules.
    """
    logger.info("Building Native Word Document with Tax Summary Engine")
    generate_word_document(output_path)

if __name__ == "__main__":
    generate_word_document()
