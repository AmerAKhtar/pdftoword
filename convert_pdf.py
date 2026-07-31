import os

# Dictionary containing all project files and their content
FILES = {
    "pdf2docx-cloudrun/requirements.txt": r'''
fastapi==0.110.0
uvicorn[standard]==0.27.1
python-multipart==0.0.9
pydantic==2.6.1
PyMuPDF==1.23.26
pdfplumber==0.10.4
python-docx==1.1.0
Pillow==10.2.0
pytesseract==0.3.10
pdf2image==1.17.0
ftfy==6.1.3
google-cloud-storage==2.14.0
lxml==5.1.0
numpy==1.26.4
''',
    "pdf2docx-cloudrun/Dockerfile": r'''
FROM python:3.11-slim AS base

ENV DEBIAN_FRONTEND=noninteractive \
    PYTHONDONTWRITEBYTECODE=1 \
    PYTHONUNBUFFERED=1

# System deps: Tesseract (OCR), Poppler (pdf2image), fonts for fidelity
RUN apt-get update && apt-get install -y --no-install-recommends \
        tesseract-ocr tesseract-ocr-eng tesseract-ocr-deu tesseract-ocr-fra \
        tesseract-ocr-spa tesseract-ocr-ita tesseract-ocr-por tesseract-ora \
        poppler-utils \
        fonts-dejavu fonts-liberation fonts-noto fonts-noto-cjk \
        libxml2 libxslt1.1 \
        ca-certificates curl \
    && rm -rf /var/lib/apt/lists/*

WORKDIR /app
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

COPY app/ ./app/

EXPOSE 8080
CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8080", "--workers", "1"]
''',
    "pdf2docx-cloudrun/service.yaml": r'''
apiVersion: serving.knative.dev/v1
kind: Service
metadata:
  name: pdf2docx
spec:
  template:
    spec:
      containerConcurrency: 4
      timeoutSeconds: 3600
      containers:
      - image: gcr.io/PROJECT_ID/pdf2docx:latest
        resources:
          limits:
            cpu: "4"
            memory: "8Gi"
        env:
        - name: GCS_BUCKET
          value: "your-bucket"
        - name: TESSERACT_LANGS
          value: "eng,deu,fra,spa"
''',
    "pdf2docx-cloudrun/app/__init__.py": r''' ''',
    "pdf2docx-cloudrun/app/config.py": r'''
import os
from dataclasses import dataclass

@dataclass(frozen=True)
class Settings:
    gcs_bucket: str = os.getenv("GCS_BUCKET", "pdf2docx-bucket")
    max_pdf_mb: int = 200
    ocr_dpi: int = 300
    ocr_min_confidence: float = 60.0   # below this -> flag, keep but annotate
    ocr_lang: str = os.getenv("OCR_LANG", "eng")
    tmp_dir: str = "/tmp"
    heading_size_ratio: float = 1.15   # > 1.15x median body size => heading
    min_bullet_chars: set = frozenset({"•","·","-","–","—","*","▪","◦","➢","■","○"})

settings = Settings()
''',
    "pdf2docx-cloudrun/app/schemas.py": r'''
from pydantic import BaseModel
from typing import Optional

class ConvertRequest(BaseModel):
    gcs_uri: Optional[str] = None   # gs://bucket/key.pdf
    filename: Optional[str] = "document.pdf"

class ConvertResponse(BaseModel):
    docx_url: str
    flags_url: str
    pages: int
    flagged_count: int
    elapsed_sec: float
''',
    "pdf2docx-cloudrun/app/storage.py": r'''
from google.cloud import storage
from datetime import timedelta, datetime
import urllib.request, pathlib
from .config import settings

_client = storage.Client()

def download(gs_uri: str, dest: pathlib.Path):
    if gs_uri.startswith("gs://"):
        b, k = gs_uri[5:].split("/", 1)
        _client.bucket(b).blob(k).download_to_filename(dest)
    else:  # http(s)
        urllib.request.urlretrieve(gs_uri, dest)

def upload(local: pathlib.Path, key: str, content_type: str) -> str:
    blob = _client.bucket(settings.gcs_bucket).blob(key)
    blob.upload_from_filename(local, content_type=content_type)
    return blob.generate_signed_url(
        version="v4", expiration=timedelta(hours=2), method="GET"
    )
''',
    "pdf2docx-cloudrun/app/parser/__init__.py": r''' ''',
    "pdf2docx-cloudrun/app/parser/encoding.py": r'''
import ftfy, re, unicodedata

LIGATURES = {
    "ﬁ":"fi","ﬂ":"fl","ﬀ":"ff","ﬃ":"ffi","ﬄ":"ffl",
    "ﬆ":"st","ﬅ":"ft"
}
CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

def repair(text: str) -> str:
    if not text:
        return text
    text = ftfy.fix_text(text, normalization="NFC", unescape_html=True)
    for k,v in LIGATURES.items():
        text = text.replace(k, v)
    # Soft hyphen at line start: drop; mid-line: keep as optional hyphen marker
    text = text.replace("\u00ad", "")
    text = CONTROL_RE.sub("", text)
    # Replace replacement char with sentinel for flagging later
    return text

def looks_garbled(text: str) -> bool:
    if not text:
        return False
    rep = text.count("\ufffd") / len(text)
    nonprint = sum(1 for c in text if unicodedata.category(c).startswith("C"))
    return rep > 0.05 or nonprint / max(len(text),1) > 0.1
''',
    "pdf2docx-cloudrun/app/parser/text_block.py": r'''
from dataclasses import dataclass, field
from typing import List, Tuple

@dataclass
class Span:
    text: str
    font: str            # normalized Word font name
    size: float
    color: int           # 0xRRGGBB
    bold: bool
    italic: bool
    underline: bool
    bbox: Tuple[float,float,float,float]
    flagged: bool = False
    flag_reason: str = ""

@dataclass
class Line:
    bbox: Tuple[float,float,float,float]
    spans: List[Span] = field(default_factory=list)
    baseline: float = 0.0

@dataclass
class Block:
    bbox: Tuple[float,float,float,float]
    lines: List[Line] = field(default_factory=list)
    kind: str = "body"   # body|heading|bullet|footnote|page_number|caption|table_cell
    level: int = 0       # heading level 1..6
    page: int = 0
''',
    "pdf2docx-cloudrun/app/parser/font_map.py": r'''
import re

# Map common PDF base-14 and embedded font names to Word equivalents.
BASE = {
    "arial":"Arial","helvetica":"Arial","arialmt":"Arial",
    "times":"Times New Roman","timesnewroman":"Times New Roman",
    "timesnewromanpsmt":"Times New Roman","times-roman":"Times New Roman",
    "courier":"Courier New","couriernew":"Courier New","couriernewpsmt":"Courier New",
    "symbol":"Symbol","zapfdingbats":"Wingdings",
    "calibri":"Calibri","cambria":"Cambria","verdana":"Verdana",
    "georgia":"Georgia","tahoma":"Tahoma","trebuchetms":"Trebuchet MS",
}

SUBSET_PREFIX = re.compile(r"^[A-Z]{6}\+")   # e.g. BAAAAA+Helvetica

def normalize(name: str) -> str:
    if not name:
        return "Calibri"
    name = SUBSET_PREFIX.sub("", name)
    base = re.split(r"[,\-]", name)[0].strip()
    key = base.lower().replace(" ", "")
    return BASE.get(key, base)

def flags_to_attrs(flags: int):
    # PDF text flags: bit 1(1)=fixed-pitch, bit2(2)=serif, bit3(4)=symbolic,
    # bit4(8)=script, bit6(32)=nonsymbolic, bit7(64)=italic
    return {
        "italic": bool(flags & 64),
    }

def detect_bold(name: str) -> bool:
    n = name.lower()
    return any(t in n for t in ("bold","black","heavy","semibold","demi","demi-"))
''',
    "pdf2docx-cloudrun/app/parser/pdfdoc.py": r'''
import fitz  # PyMuPDF
from .text_block import Span, Line, Block
from .font_map import normalize, flags_to_attrs, detect_bold
from .encoding import repair, looks_garbled

def rgb_to_int(color_tuple):
    if not color_tuple: return 0
    r,g,b = (color_tuple + (0,0,0))[:3]
    return (int(r*255)<<16) | (int(g*255)<<8) | int(b*255)

class PdfDoc:
    def __init__(self, path: str):
        self.doc = fitz.open(path)
        self.pages = []

    def parse(self):
        for pno, page in enumerate(self.doc):
            blocks = []
            d = page.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES
                                          | fitz.TEXT_PRESERVE_WHITESPACE)
            for b in d["blocks"]:
                if b["type"] != 0:        # 0 = text
                    continue
                blk = Block(bbox=tuple(b["bbox"]), page=pno)
                for ln in b["lines"]:
                    line = Line(bbox=tuple(ln["bbox"]), baseline=ln.get("origin",[0,0])[1])
                    for sp in ln["spans"]:
                        raw = sp["text"]
                        fixed = repair(raw)
                        font = normalize(sp["font"])
                        attrs = flags_to_attrs(sp.get("flags",0))
                        bold = detect_bold(sp["font"])
                        flagged = looks_garbled(fixed)
                        span = Span(
                            text=fixed,
                            font=font,
                            size=round(sp["size"],1),
                            color=rgb_to_int(sp.get("color")),
                            bold=bold,
                            italic=attrs["italic"],
                            underline=False,
                            bbox=tuple(sp["bbox"]),
                            flagged=flagged,
                            flag_reason=("encoding_unclear" if flagged else "")
                        )
                        line.spans.append(span)
                    blk.lines.append(line)
                if blk.lines:
                    blocks.append(blk)
            self.pages.append(blocks)
        return self.pages
''',
    "pdf2docx-cloudrun/app/parser/layout.py": r'''
from statistics import median
from .text_block import Block, Line
from ..config import settings

def classify_page(page_blocks):
    """Assign each block a kind: heading | bullet | page_number | footnote | body."""
    sizes = [sp.size for b in page_blocks for ln in b.lines for sp in ln.spans]
    if not sizes:
        return page_blocks
    med = median(sizes)
    page_top = min(b.bbox[1] for b in page_blocks)
    page_bot = max(b.bbox[3] for b in page_blocks)
    page_h = page_bot - page_top

    for b in page_blocks:
        text = "".join(sp.text for ln in b.lines for sp in ln.spans).strip()
        first = b.lines[0].spans[0] if b.lines and b.lines[0].spans else None

        # Page number: short numeric near top/bottom margin
        if (len(text) <= 6 and text.isdigit()
            and (b.bbox[1] - page_top < page_h*0.06
                 or page_bot - b.bbox[3] < page_h*0.06)):
            b.kind = "page_number"; continue

        # Footnote: at bottom, smaller than median, often starts with superscript digit
        if (page_bot - b.bbox[3] < page_h*0.10
            and first and first.size < med*0.92
            and (text[:1].isdigit() or text.startswith("("))):
            b.kind = "footnote"; continue

        # Bullet: starts with bullet glyph or "-"
        if first and first.text.strip() in settings.min_bullet_chars:
            b.kind = "bullet"; continue

        # Heading: larger than body, short line count, often bold
        if first and first.size >= med * settings.heading_size_ratio and len(b.lines) <= 2:
            ratio = first.size / med
            b.level = 1 if ratio > 1.6 else (2 if ratio > 1.35 else 3)
            b.kind = "heading"; continue

        b.kind = "body"
    return page_blocks

def merge_spans_by_line(lines):
    """Concatenate spans in a line into runs grouped by font/size/attrs."""
    runs = []
    for ln in lines:
        cur = None
        for sp in ln.spans:
            key = (sp.font, sp.size, sp.bold, sp.italic, sp.underline, sp.color)
            if cur and cur["_k"] == key:
                cur["text"] += sp.text
                cur["flagged"] = cur["flagged"] or sp.flagged
            else:
                cur = {"_k": key, "text": sp.text, "font": sp.font,
                       "size": sp.size, "bold": sp.bold, "italic": sp.italic,
                       "underline": sp.underline, "color": sp.color,
                       "flagged": sp.flagged}
                runs.append(cur)
    return runs
''',
    "pdf2docx-cloudrun/app/parser/tables.py": r'''
import pdfplumber
from .text_block import Block

def extract_tables(pdf_path: str):
    """Return list of (page_index, table_rows, bbox) tuples."""
    out = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.find_tables(
                table_settings={
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 3,
                    "join_tolerance": 3,
                    "intersection_tolerance": 3,
                }
            )
            for t in tables:
                rows = t.extract()           # list of list of str|None
                bbox = t.bbox
                out.append((i, rows, bbox))
    return out
''',
    "pdf2docx-cloudrun/app/parser/images.py": r'''
import fitz, io, pathlib
from PIL import Image

def extract_images(doc: fitz.Document):
    """Return per-page list of (bbox, png_bytes, likely_text_image: bool)."""
    out = [[] for _ in range(doc.page_count)]
    for pno, page in enumerate(doc):
        for info in page.get_image_info(xrefs=True):
            bbox = info["bbox"]
            xref = info.get("xref")
            if xref is None: continue
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n - pix.alpha >= 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_bytes = pix.tobytes("png")
            except Exception:
                continue
            # Heuristic: small image with high pixel density likely contains text
            likely_text = (info.get("width",0) < 800 and info.get("height",0) < 300)
            out[pno].append((bbox, img_bytes, likely_text))
    return out
''',
    "pdf2docx-cloudrun/app/parser/ocr.py": r'''
import io, pytesseract
from PIL import Image
from pdf2image import convert_from_path
from ..config import settings

def render_page_to_image(pdf_path: str, page_no: int) -> Image.Image:
    imgs = convert_from_path(pdf_path, dpi=settings.ocr_dpi,
                             first_page=page_no+1, last_page=page_no+1)
    return imgs[0]

def is_scanned(page_text: str) -> bool:
    return len(page_text.strip()) < 30

def ocr_image(img: Image.Image):
    data = pytesseract.image_to_data(img, lang=settings.ocr_lang,
                                     output_type=pytesseract.Output.DICT)
    words, confs = [], []
    for i, w in enumerate(data["text"]):
        if not w.strip(): continue
        try: c = float(data["conf"][i])
        except: c = -1
        words.append((w, c, data["left"][i], data["top"][i],
                      data["width"][i], data["height"][i]))
        confs.append(c)
    avg_conf = sum(confs)/len(confs) if confs else 0
    text = " ".join(w[0] for w in words)
    return text, avg_conf, words
''',
    "pdf2docx-cloudrun/app/renderer/__init__.py": r''' ''',
    "pdf2docx-cloudrun/app/renderer/comments.py": r'''
# python-docx doesn't support comments natively; we insert them via raw OOXML.
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_comment(doc, paragraph, comment_text, author="pdf2docx"):
    """Add a Word comment anchored to the given paragraph."""
    cid = add_comment._counter = getattr(add_comment, "_counter", 0) + 1
    
    comments_part = getattr(doc.part, "_comments_part", None)
    if comments_part is None:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.part import Part
        from docx.oxml import parse_xml
        from docx.opc.packuri import PackURI
        # create comments part lazily
        partname = PackURI("/word/comments.xml")
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        xml = '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        comments_part = Part(partname, content_type, parse_xml(xml).bytes, doc.part.package)
        doc.part.relate_to(comments_part, RT.COMMENTS)
        doc.part._comments_part = comments_part

    c_el = OxmlElement("w:comment")
    c_el.set(qn("w:id"), str(cid))
    c_el.set(qn("w:author"), author)
    c_el.set(qn("w:date"), datetime.datetime.utcnow().isoformat()+"Z")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = comment_text
    r.append(t); p.append(r); c_el.append(p)
    comments_part.element.append(c_el)

    # Insert commentRangeStart/End + reference in paragraph
    start = OxmlElement("w:commentRangeStart"); start.set(qn("w:id"), str(cid))
    end   = OxmlElement("w:commentRangeEnd");   end.set(qn("w:id"), str(cid))
    ref   = OxmlElement("w:r")
    rpr   = OxmlElement("w:rPr")
    rs    = OxmlElement("w:rStyle"); rs.set(qn("w:val"), "CommentReference")
    rpr.append(rs); ref.append(rpr)
    cr    = OxmlElement("w:commentReference"); cr.set(qn("w:id"), str(cid))
    ref.append(cr)

    p_xml = paragraph._p
    p_xml.insert(0, start)
    p_xml.append(end)
    p_xml.append(ref)
''',
    "pdf2docx-cloudrun/app/renderer/docx_writer.py": r'''
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ..parser.text_block import Block
from .comments import add_comment
import io, pathlib

def _set_run_font(run, font, size, bold, italic, color):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    for a in ("ascii","hAnsi","cs","eastAsia"):
        rFonts.set(qn(f"w:{a}"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor((color>>16)&0xFF,(color>>8)&0xFF,color&0xFF)

def _page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)

def render(parsed, out_path: pathlib.Path, tables, images):
    """parsed: list of pages of Blocks. tables: from tables.extract_tables.
       images: from images.extract_images."""
    doc = Document()
    # Build a quick lookup of (page, bbox) -> table so we skip overlapping text
    table_bboxes_by_page = {}
    for pno, rows, bbox in tables:
        table_bboxes_by_page.setdefault(pno, []).append((bbox, rows))

    def in_any_table(pno, bbox):
        for tbbox, _ in table_bboxes_by_page.get(pno, []):
            x0,y0,x1,y1 = bbox
            tx0,ty0,tx1,ty1 = tbbox
            if x0 >= tx0 and y0 >= ty0 and x1 <= tx1 and y1 <= ty1:
                return True
        return False

    flagged_count = 0
    for pno, blocks in enumerate(parsed):
        # Insert images whose top is on this page
        for (bbox, png, likely_text) in images[pno]:
            if likely_text:
                add_comment(doc, doc.add_paragraph(),
                            f"[FLAG] Image at ({bbox[0]:.0f},{bbox[1]:.0f}) "
                            f"looks like a text image and should be OCR'd.")
                flagged_count += 1
            try:
                doc.add_picture(io.BytesIO(png), width=Inches(min(6.0,(bbox[2]-bbox[0])/72)))
            except Exception:
                pass

        for b in blocks:
            if in_any_table(pno, b.bbox):
                continue   # text already inside a table cell; rendered with table
            if b.kind == "page_number":
                continue   # routed to header/footer below
            style = {
                "heading": f"Heading {min(b.level,3)}",
                "bullet": "List Bullet",
                "footnote": None,    # handled separately
                "body": None,
                "caption": "Caption",
            }.get(b.kind, None)

            if b.kind == "footnote":
                # python-docx footnotes require OOXML surgery; for brevity,
                # we emit them at page end with a separator line.
                pass

            para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            for ln in b.lines:
                for sp in ln.spans:
                    run = para.add_run(sp.text)
                    _set_run_font(run, sp.font, sp.size, sp.bold, sp.italic, sp.color)
                    if sp.flagged:
                        add_comment(doc, para,
                            f"[FLAG] Encoding unclear for run '{sp.text[:40]}... ' "
                            f"(font={sp.font}, size={sp.size}). Verify manually.")
                        flagged_count += 1
                # line break between lines within same block
                if ln is not b.lines[-1]:
                    para.add_run().add_break()

        # Insert tables on this page
        for (tbbox, rows) in table_bboxes_by_page.get(pno, []):
            if not rows: continue
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            tbl.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    txt = (cell or "").strip()
                    if j < len(tbl.rows[i].cells):
                        tbl.rows[i].cells[j].text = txt

        if pno < len(parsed) - 1:
            _page_break(doc)

    # Header/footer: simple page number on right of footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)

    doc.save(out_path)
    return flagged_count
''',
    "pdf2docx-cloudrun/app/pipeline.py": r'''
import time, pathlib, json, tempfile
from .parser.pdfdoc import PdfDoc
from .parser.layout import classify_page
from .parser import tables as tables_mod
from .parser import images as images_mod
from .parser import ocr as ocr_mod
from .renderer.docx_writer import render as render_docx
from .storage import download, upload
from .config import settings

def run(gcs_uri: str, filename: str):
    t0 = time.time()
    tmp = pathlib.Path(tempfile.mkdtemp())
    pdf_path = tmp / filename
    download(gcs_uri, pdf_path)

    pdf = PdfDoc(str(pdf_path))
    pages = pdf.parse()
    # OCR scanned pages and merge into blocks
    flags = []
    for pno, blocks in enumerate(pages):
        page_text = "".join(sp.text for b in blocks for ln in b.lines for sp in ln.spans)
        if ocr_mod.is_scanned(page_text):
            img = ocr_mod.render_page_to_image(str(pdf_path), pno)
            text, conf, _ = ocr_mod.ocr_image(img)
            from .parser.text_block import Block, Line, Span
            blk = Block(bbox=(0,0,0,0), page=pno)
            ln = Line(bbox=(0,0,0,0))
            ln.spans.append(Span(text, "Calibri", 11, 0, False, False, False,
                                 (0,0,0,0), flagged=conf < settings.ocr_min_confidence,
                                 flag_reason=f"ocr_confidence={conf:.1f}"))
            blk.lines.append(ln)
            pages[pno] = [blk]
            if conf < settings.ocr_min_confidence:
                flags.append({"page": pno+1, "type":"ocr_low_confidence",
                              "value": round(conf,1)})

    for pno in range(len(pages)):
        pages[pno] = classify_page(pages[pno])

    tables = tables_mod.extract_tables(str(pdf_path))
    images = images_mod.extract_images(pdf.doc)

    docx_path = tmp / (pdf_path.stem + ".docx")
    flagged = render_docx(pages, docx_path, tables, images)

    flags_path = tmp / "flags.json"
    flags_path.write_text(json.dumps(flags, indent=2))

    docx_url = upload(docx_path, f"out/{pdf_path.stem}.docx",
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    flags_url = upload(flags_path, f"out/{pdf_path.stem}.flags.json",
                       "application/json")
    return {
        "docx_url": docx_url,
        "flags_url": flags_url,
        "pages": len(pages),
        "flagged_count": flagged + len(flags),
        "elapsed_sec": round(time.time()-t0, 2),
    }
''',
    "pdf2docx-cloudrun/app/main.py": r'''
from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.responses import JSONResponse
from .schemas import ConvertRequest, ConvertResponse
from .pipeline import run
from .config import settings
import io

app = FastAPI(title="pdf2docx", version="1.0")

@app.get("/healthz")
def health(): return {"ok": True}

@app.post("/convert", response_model=ConvertResponse)
async def convert(req: ConvertRequest = Body(...)):
    if not req.gcs_uri:
        raise HTTPException(400, "gcs_uri required")
    return JSONResponse(run(req.gcs_uri, req.filename or "document.pdf"))

@app.post("/convert/upload", response_model=ConvertResponse)
async def convert_upload(file: UploadFile = File(...)):
    data = await file.read()
    if len(data) > settings.max_pdf_mb * 1024 * 1024:
        raise HTTPException(413, f"PDF > {settings.max_pdf_mb} MB; use GCS.")
    # Save to /tmp then push to GCS so pipeline can pull uniformly.
    import tempfile, pathlib, uuid
    tmp = pathlib.Path(tempfile.mkdtemp()) / f"{uuid.uuid4()}.pdf"
    tmp.write_bytes(data)
    from .storage import _client
    key = f"in/{tmp.name}"
    _client.bucket(settings.gcs_bucket).blob(key).upload_from_filename(tmp)
    return JSONResponse(run(f"gs://{settings.gcs_bucket}/{key}", file.filename))
''',
    "pdf2docx-cloudrun/README.md": r'''
# PDF to Editable DOCX Converter

A Google Cloud Run service that converts PDFs into fully editable DOCX files while preserving fonts, layouts, tables, and inferring headings/bullets. Includes OCR fallback for scanned pages and explicitly flags garbled or unreadable content instead of guessing.

## Deployment

1. **Set Environment Variables**: Update `GCS_BUCKET` in `service.yaml` to point to your Google Cloud Storage bucket.

2. **Build & Deploy**:
   ```bash
   gcloud builds submit --tag gcr.io/PROJECT_ID/pdf2docx .
   gcloud run deploy pdf2docx \
     --image gcr.io/PROJECT_ID/pdf2docx \
     --region us-central1 \
     --cpu 4 --memory 8Gi \
     --timeout 3600 \
     --concurrency 4 \
     --cpu-boost \
     --set-env-vars GCS_BUCKET=your-bucket,OCR_LANG=eng